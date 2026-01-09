import os
import pycountry
import pandas as pd
from flask import Flask, render_template, request, jsonify, send_file
from google import genai
from pypdf import PdfReader
from docx import Document
from PIL import Image

app = Flask(__name__)

# SECURITY NOTE: This key is currently hitting quota limits.
# Consider generating a new one at aistudio.google.com
GEMINI_API_KEY = "AIzaSyBfg4MyB9EuZVphGN2sSviZJ8PTwBV1z4o"
client = genai.Client(api_key=GEMINI_API_KEY)

# Ensure upload folder exists
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# --- Helper Functions ---

def get_languages():
    """Returns a sorted list of languages for the dropdowns."""
    return sorted([
        {'code': lang.alpha_2, 'name': lang.name}
        for lang in pycountry.languages if hasattr(lang, 'alpha_2')
    ], key=lambda x: x['name'])


# --- Conversion Logic ---

def convert_pdf_to_docx(input_path):
    output_path = input_path.replace('.pdf', '.docx')
    reader = PdfReader(input_path)
    doc = Document()
    for page in reader.pages:
        text = page.extract_text()
        if text:
            doc.add_paragraph(text)
    doc.save(output_path)
    return output_path


def extract_word_to_excel(input_path):
    output_path = input_path.replace('.docx', '.xlsx')
    doc = Document(input_path)
    all_data = []
    for table in doc.tables:
        for row in table.rows:
            all_data.append([cell.text.strip() for cell in row.cells])
    if all_data:
        df = pd.DataFrame(all_data)
        df.to_excel(output_path, index=False, header=False)
        return output_path
    return None


def image_to_pdf(input_path):
    output_path = input_path.rsplit('.', 1)[0] + '.pdf'
    image = Image.open(input_path)
    if image.mode == 'RGBA':
        image = image.convert('RGB')
    image.save(output_path, 'PDF', resolution=100.0)
    return output_path


# --- Routes ---

@app.route('/')
def index():
    return render_template('index.html', languages=get_languages())


@app.route('/translator')
def translator():
    return render_template('translator.html', languages=get_languages())


@app.route('/converter')
def converter():
    return render_template('file_convertor.html')


@app.route('/convert', methods=['POST'])
def convert():
    if 'file' not in request.files:
        return "No file uploaded", 400

    file = request.files['file']
    # FIX: Get the target format from the form selection
    target_format = request.form.get('target_format')

    if file.filename == '':
        return "No selected file", 400

    file_path = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(file_path)
    extension = file.filename.split('.')[-1].lower()

    try:
        output_path = None

        # FIX: Added function calls () and variable validation
        if extension == 'pdf' and target_format == 'docx':
            output_path = convert_pdf_to_docx(file_path)
        elif extension == 'docx' and target_format == 'xlsx':
            output_path = extract_word_to_excel(file_path)
        elif extension in ['jpg', 'jpeg', 'png'] and target_format == 'pdf':
            output_path = image_to_pdf(file_path)
        else:
            return f"Unsupported conversion: {extension} to {target_format}", 400

        if output_path and os.path.exists(output_path):
            return send_file(output_path, as_attachment=True)
        return "Conversion Failed: Output file not generated.", 500

    except Exception as e:
        return f"Error: {str(e)}", 500


@app.route('/translate', methods=['POST'])
def translate():
    data = request.get_json()
    target_lang = data.get('target_lang')
    text = data.get('text')

    if not text or not target_lang:
        return jsonify({"error": "Missing input text"}), 400

    prompt = f"Translate the following text to {target_lang}. Output ONLY the translation: {text}"

    try:
        response = client.models.generate_content(
            model='gemini-2.0-flash',
            contents=prompt
        )
        return jsonify({"translated": response.text.strip()})
    except Exception as e:
        # Check for Quota/Rate Limit specifically
        if "429" in str(e):
            return jsonify({"error": "AI Quota exceeded. Please wait 30 seconds and try again."}), 429
        return jsonify({"error": str(e)}), 500


if __name__ == '__main__':
    app.run(debug=True)